from superpowered.main import BASE_URL, _format_http_response, get_headers
import requests
import os
import PyPDF2
import docx
from bs4 import BeautifulSoup

class KnowledgeBase:
    """
    The KnowledgeBase class is used to create, update, and delete knowledge bases. It can also be used to directly query a knowledge base.
    """
    def __init__(self, title: str, supp_id: str = None, description: str = None, kb_id: str = None):
        self.title = title
        self.supp_id = supp_id
        self.description = description
        self.kb_id = kb_id
        self.is_deployed = self.kb_id is not None

    def create(self):
        if self.is_deployed:
            raise Exception('This knowledge base has already been deployed: ' + self.kb_id)
        url = BASE_URL + 'knowledge_bases'
        payload = {
            'title': self.title
        }
        if self.supp_id is not None:
            payload['supp_id'] = self.supp_id
        if self.description is not None:
            payload['description'] = self.description
        resp = _format_http_response(requests.post(url, headers=get_headers(), json=payload))
        self.kb_id = resp['body']['id']
        self.is_deployed = True
        return resp['body']

    def add_document(self, content: str, title: str = None, link_to_source: str = None, supp_id: str = None, description: str = None):
        kb_document = KnowledgeBaseDocument(kb_id=self.kb_id, content=content, title=title, link_to_source=link_to_source, supp_id=supp_id, description=description)
        resp_body = kb_document.create()
        return resp_body

    def get_documents(self):
        url = BASE_URL + f'knowledge_bases/{self.kb_id}/documents'
        resp = _format_http_response(requests.get(url, headers=get_headers()))
        documents = {}
        for document in resp['body']['documents']:
            kb_document = KnowledgeBaseDocument(
                kb_id=document['id'],
                content=None,
                title=document['title'],
                link_to_source=document['link_to_source'],
                supp_id=document['supp_id'],
                description=document['description'],
            )
            documents[document['id']] = kb_document
        return documents

    def query(self, query: str, retriever_top_k: int = 100, reranker_top_k: int = 5, extract_and_summarize: bool = False):
        """
        directly query a knowledge base
        """
        url = BASE_URL + f'knowledge_bases/query'
        payload = {
            'query': query,
            'retriever_top_k': retriever_top_k,
            'reranker_top_k': reranker_top_k,
            'knowledge_base_ids': [self.kb_id],
            'summarize_results': extract_and_summarize,
        }
        resp = _format_http_response(requests.post(url, headers=get_headers(), json=payload))
        return resp["body"]

    def delete(self):
        url = BASE_URL + f'knowledge_bases/{self.kb_id}'
        resp = _format_http_response(requests.delete(url, headers=get_headers()))
        return resp


class KnowledgeBaseDocument:
    """
    The KnowledgeBaseDocument class is used to create and delete documents in a knowledge base.
    """
    def __init__(self, kb_id: str, content: str = None, title: str = None, link_to_source: str = None, supp_id: str = None, description: str = None, doc_id: str = None):
        self.kb_id = kb_id
        self.content = content # content is a string
        self.title = title
        self.link_to_source = link_to_source
        self.supp_id = supp_id
        self.description = description
        self.doc_id = doc_id

    def create(self):
        url = BASE_URL + f'knowledge_bases/{self.kb_id}/documents/raw_text'
        payload = {
            'content': self.content
        }
        if self.title is not None:
            payload['title'] = self.title
        if self.link_to_source is not None:
            payload['link_to_source'] = self.link_to_source
        if self.supp_id is not None:
            payload['supp_id'] = self.supp_id
        if self.description is not None:
            payload['description'] = self.description
        resp = _format_http_response(requests.post(url, headers=get_headers(), json=payload))
        return resp['body']

    def delete(self):
        url = BASE_URL + f'knowledge_bases/{self.kb_id}/documents/{self.doc_id}'
        resp = _format_http_response(requests.delete(url, headers=get_headers()))
        return resp


def get_knowledge_bases():
    knowledge_bases = {}
    url = BASE_URL + 'knowledge_bases'
    resp = requests.get(url, headers=get_headers())
    resp = _format_http_response(resp)
    for knowledge_base in resp['body']['knowledge_bases']:
        kb = KnowledgeBase(
            title=knowledge_base['title'],
            kb_id=knowledge_base['id'],
            supp_id = knowledge_base['supp_id'],
            description=knowledge_base['description']
        )
        knowledge_bases[knowledge_base['id']] = kb
    return knowledge_bases

def create_knowledge_base(title: str, supp_id: str = None, description: str = None):
    """
    create_knowledge_base() is a convenience function that creates a KnowledgeBase object and then calls its create() method
    """
    kb = KnowledgeBase(title, supp_id, description)
    resp = kb.create()
    return kb

def get_knowledge_base(title: str):
    """
    get_knowledge_base() is a convenience function that returns a KnowledgeBase object for an existing knowledge base, given its title
    """
    knowledge_bases = get_knowledge_bases()
    # create knowledge base name to id map - knowledge_bases is a dict of knowledge base objects keyed on knowledge_base_id
    kb_from_title = {}
    for kb_id in knowledge_bases.keys():
        kb = knowledge_bases[kb_id] # get the KnowledgeBase object
        if kb.title not in kb_from_title:
            kb_from_title[kb.title] = kb
        else:
            print ('WARNING: Duplicate knowledge base title: ' + kb.title)
    
    # if the title we're looking for is in the map, return the KnowledgeBase object
    if title in kb_from_title:
        return kb_from_title[title]
    else:
        raise Exception('Knowledge base title not found: ' + title)

def list_knowledge_bases(verbose=True):
    """
    list_knowledge_bases() is a convenience function that returns a dictionary of all KnowledgeBase objects for an account
    """
    knowledge_bases = get_knowledge_bases()
    if verbose:
        print ("\nKnowledge bases:")
        for kb_obj in knowledge_bases.values():
            print (f"id: {kb_obj.kb_id}\ntitle: {kb_obj.title}\n")
    return knowledge_bases

def list_documents_in_kb(kb_title: str, verbose=True):
    """
    list_documents_in_kb() is a convenience function that returns a dictionary of KnowledgeBaseDocument objects for an existing knowledge base, given its title
    """
    kb = get_knowledge_base(kb_title)
    documents = kb.get_documents()
    if verbose:
        print ("\nDocuments:")
        for document_obj in documents.values():
            print (f"title: {document_obj.title}\n\n")
    return documents

def read_file(file_path: str):
    """
    read a file and return its contents as a string
    """
    # infer the file type from the file extension
    file_type = file_path.split(".")[-1]

    if file_type == "pdf":
        # Open the PDF file
        with open(file_path, 'rb') as file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(file)

            # Initialize an empty string to store the extracted text
            extracted_text = ""

            # Loop through all the pages in the PDF file
            for page_num in range(len(pdf_reader.pages)):
                # Extract the text from the current page
                page_text = pdf_reader.pages[page_num].extract_text()

                # Add the extracted text to the final text
                extracted_text += page_text 
        return extracted_text
    elif file_type == "txt":
        with open(file_path, 'r') as f:
            # read the file and convert it to a string
            content = f.read()
        return content
    elif file_type == "docx":
        # Create a docx Document object
        doc = docx.Document(file_path)

        # Initialize an empty string to store the extracted text
        extracted_text = ""

        # Loop through all the paragraphs in the docx file
        for paragraph in doc.paragraphs:
            # Add the extracted text to the final text
            extracted_text += paragraph.text + "\n"

        return extracted_text
    elif file_type == "md":
        with open(file_path, 'r') as f:
            # read the file and convert it to a string
            content = f.read()
        return content
    else:
        raise Exception("File type not supported. Supported file types are: .txt, .pdf, .docx, and .md")
    
def add_document_to_kb(kb_title: str, content: str, title: str = None, link_to_source: str = None, supp_id: str = None, description: str = None):
    """
    add a document to a knowledge base by providing the title and content (as a string)
    - use this function if you want to directly add text to a knowledge base. If you want to add a file, use add_file_to_kb() instead.
    """
    kb = get_knowledge_base(kb_title)
    kb.add_document(content=content, title=title, link_to_source=link_to_source, supp_id=supp_id, description=description)

def add_file_to_kb(kb_title: str, file_path: str, title: str = None, link_to_source: str = None, supp_id: str = None, description: str = None):
    """
    add a Document to a knowledge base by reading the contents of a file - currently supports .txt, .pdf, .docx, and .md files
    """
    content = read_file(file_path)
    
    # make the title the file name if it's not provided
    if title is None:
        title = file_path.split("/")[-1]
    
    kb = get_knowledge_base(kb_title)
    kb.add_document(content=content, title=title, link_to_source=link_to_source, supp_id=supp_id, description=description)

def add_directory_to_kb(kb_title: str, directory_path: str, verbose=False):
    """
    add all the supported files in a directory to a knowledge base
    """
    # get a list of all the files in the directory
    all_files = []
    for root, _, files in os.walk(directory_path):
        for file in files:
            if file.endswith(".txt") or file.endswith(".pdf") or file.endswith(".docx") or file.endswith(".md"):
                all_files.append(os.path.join(root, file))

    # add each file to the knowledge base
    for file in all_files:
        if verbose: print("Adding:", file)
        add_file_to_kb(kb_title, file_path=file)

def add_url_to_kb(kb_title: str, url: str, title: str = None, link_to_source: str = None, supp_id: str = None, description: str = None):
    """
    add a Document to a knowledge base by scraping the contents of a URL
    """
    # scrape the URL
    response = requests.get(url)
    soup = BeautifulSoup(response.content, "html.parser")
    body = soup.find('body')
    content = body.text

    # make the title the URL if no title is provided
    if title is None:
        title = url

    kb = get_knowledge_base(kb_title)
    kb.add_document(content=content, title=title, link_to_source=link_to_source, supp_id=supp_id, description=description)

def query(query: str, kb_titles: list = [], kb_ids: list = [], retriever_top_k: int = 100, reranker_top_k: int = 5, extract_and_summarize: bool = False):
    """
    Can take either a list of kb_titles or a list of kb_ids (or a mix of both)
    """
    # convert kb_titles to kb_ids
    if len(kb_titles) > 0:
        knowledge_bases = get_knowledge_bases() # gets a dict of KnowledgeBase objects keyed on kb_id
        for kb_title in kb_titles:
            # see if the title is in the knowledge bases dict
            for kb_id in knowledge_bases.keys():
                if knowledge_bases[kb_id].title == kb_title:
                    kb_ids.append(kb_id)
                    break

    # deduplicate kb_ids
    kb_ids = list(set(kb_ids))

    # make sure kb_ids is not empty
    if len(kb_ids) == 0:
        raise Exception("Must provide at least one valid knowledge base title or id")

    # make the request
    url = BASE_URL + f'knowledge_bases/query'
    payload = {
        'query': query,
        'knowledge_base_ids': kb_ids,
        'retriever_top_k': retriever_top_k,
        'reranker_top_k': reranker_top_k,
        'summarize_results': extract_and_summarize,
    }
    resp = _format_http_response(requests.post(url, headers=get_headers(), json=payload))
    return resp["body"]